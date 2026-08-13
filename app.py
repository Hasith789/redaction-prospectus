import streamlit as st
from pathlib import Path
import sys
import tempfile

sys.path.insert(0, "src")

from detector import detect_pii

st.set_page_config(
    page_title="PII Redaction Tool",
    page_icon="🔒",
    layout="centered"
)

st.title("🔒 PII Redaction Tool")
st.write("Upload a DOCX document to detect and redact personally identifiable information.")

uploaded_file = st.file_uploader(
    "Upload DOCX file",
    type=["docx"]
)

if uploaded_file:

    st.success(f"Uploaded: {uploaded_file.name}")

    if st.button("Redact Document", type="primary"):

        with tempfile.NamedTemporaryFile(
            delete=False,
            suffix=".docx"
        ) as temp:

            temp.write(uploaded_file.getbuffer())
            input_path = Path(temp.name)

        st.info("Processing document...")

        # Import your existing redaction functions
        from docx import Document
        from redact_document import redact_text, process_table

        document = Document(input_path)

        for paragraph in document.paragraphs:
            paragraph.text = redact_text(paragraph.text)

        for table in document.tables:
            process_table(table)

        output_path = Path(tempfile.mktemp(suffix="_redacted.docx"))

        document.save(output_path)

        st.success("✅ Redaction completed!")

        with open(output_path, "rb") as f:
            st.download_button(
                label="⬇️ Download Redacted DOCX",
                data=f,
                file_name="redacted_prospectus.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )