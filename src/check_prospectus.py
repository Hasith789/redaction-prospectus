from pathlib import Path

from docx import Document

from detector import detect_pii


INPUT_FILE = Path("input/Red Herring Prospectus.docx")


def extract_text(docx_path):
    document = Document(docx_path)

    parts = []

    # Paragraphs
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    # Tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n".join(parts)


def main():
    print("Reading:", INPUT_FILE)

    text = extract_text(INPUT_FILE)

    print(f"Extracted characters: {len(text):,}")

    detections = detect_pii(text)

    print(f"Total Stage-1 detections: {len(detections)}")
    print()

    counts = {}

    for detection in detections:
        counts[detection.pii_type] = (
            counts.get(detection.pii_type, 0) + 1
        )

    print("Detection summary")
    print("-----------------")

    for pii_type, count in sorted(counts.items()):
        print(f"{pii_type:20} {count}")

    print()
    print("Sample detections")
    print("------------------")

    for detection in detections[:30]:
        print(
            f"[{detection.pii_type}] "
            f"{detection.value}"
        )


if __name__ == "__main__":
    main()