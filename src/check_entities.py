from pathlib import Path

from docx import Document

from entity_detector import (
    detect_people,
    detect_companies,
    detect_addresses,
)


INPUT_FILE = Path("input/Red Herring Prospectus.docx")


def extract_text(docx_path):
    document = Document(docx_path)

    parts = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n".join(parts)


def main():

    text = extract_text(INPUT_FILE)

    print("=" * 60)
    print("PERSON DETECTIONS")
    print("=" * 60)

    people = detect_people(text)

    seen = set()

    for match in people:

        if match.value not in seen:
            print(match.value)
            seen.add(match.value)

    print()
    print("Unique people:", len(seen))

    print()
    print("=" * 60)
    print("COMPANY DETECTIONS")
    print("=" * 60)

    companies = detect_companies(text)

    seen = set()

    for match in companies:

        if match.value not in seen:
            print(match.value)
            seen.add(match.value)

    print()
    print("Unique companies:", len(seen))

    print()
    print("=" * 60)
    print("ADDRESS DETECTIONS")
    print("=" * 60)

    addresses = detect_addresses(text)

    for match in addresses[:30]:
        print("-" * 40)
        print(match.value)

    print()
    print("Address detections:", len(addresses))


if __name__ == "__main__":
    main()
    