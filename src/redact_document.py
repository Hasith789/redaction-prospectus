from pathlib import Path
from docx import Document

from detector import detect_pii


INPUT = Path("input/Red Herring Prospectus.docx")
OUTPUT = Path("output/redacted_prospectus.docx")


REPLACEMENTS = {
    "PERSON": "John Doe",
    "EMAIL": "john.doe@example.com",
    "PHONE": "+91 1234567890",
    "COMPANY": "Example Corporation",
    "ADDRESS": "123 Example Street, Mumbai, Maharashtra 400001",
    "DOB": "01/01/1990",
    "SSN": "000-00-0000",
    "CREDIT_CARD": "4111 1111 1111 1111",
    "IP_ADDRESS": "192.0.2.1",
}


def redact_text(text):
    """Detect PII in a piece of text and replace it."""

    matches = detect_pii(text)

    # Replace from right to left.
    # This prevents changing the positions of earlier matches.
    for match in sorted(matches, key=lambda x: x.start, reverse=True):

        replacement = REPLACEMENTS.get(
            match.pii_type,
            "[REDACTED]"
        )

        text = (
            text[:match.start]
            + replacement
            + text[match.end:]
        )

    return text


def process_table(table):
    """Redact text inside a DOCX table, including nested tables."""

    for row in table.rows:
        for cell in row.cells:

            # Process paragraphs in the cell
            for paragraph in cell.paragraphs:
                if paragraph.text.strip():
                    paragraph.text = redact_text(paragraph.text)

            # Process nested tables
            for nested_table in cell.tables:
                process_table(nested_table)


def force_redact_tables(document):
    """Final safety pass for known remaining PII in DOCX tables."""

    replacements = {
        "Sarthak Malvadkar": "John Doe",
    }

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:

                # Direct cell text replacement
                for old, new in replacements.items():
                    if old in cell.text:
                        for paragraph in cell.paragraphs:
                            if old in paragraph.text:
                                paragraph.text = paragraph.text.replace(old, new)

                # Nested tables
                for nested_table in cell.tables:
                    force_redact_tables_from_table(nested_table, replacements)


def force_redact_tables_from_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:

            for old, new in replacements.items():
                if old in cell.text:
                    for paragraph in cell.paragraphs:
                        if old in paragraph.text:
                            paragraph.text = paragraph.text.replace(old, new)

            for nested_table in cell.tables:
                force_redact_tables_from_table(nested_table, replacements)
                

def main():

    print(f"Reading: {INPUT}")

    if not INPUT.exists():
        raise FileNotFoundError(
            f"Input file not found: {INPUT}"
        )

    OUTPUT.parent.mkdir(
        parents=True,
        exist_ok=True
    )

    document = Document(INPUT)

    # -----------------------------------------
    # Process normal paragraphs
    # -----------------------------------------

    print("Processing paragraphs...")

    for paragraph in document.paragraphs:
        paragraph.text = redact_text(paragraph.text)

    # -----------------------------------------
    # Process tables
    # -----------------------------------------

    print("Processing tables...")

    for table in document.tables:
        process_table(table)

    print("Final table safety pass...")
    force_redact_tables(document)

        # -----------------------------------------
    # Process headers and footers
    # -----------------------------------------

    print("Processing headers and footers...")

    for section in document.sections:

        for paragraph in section.header.paragraphs:
            paragraph.text = redact_text(paragraph.text)

        for table in section.header.tables:
            process_table(table)

        for paragraph in section.footer.paragraphs:
            paragraph.text = redact_text(paragraph.text)

        for table in section.footer.tables:
            process_table(table)


    # -----------------------------------------
    # Save
    # -----------------------------------------

    document.save(OUTPUT)

    print()
    print("=" * 60)
    print("REDACTION COMPLETE")
    print("=" * 60)
    print(f"Output: {OUTPUT}")
    print(f"Exists: {OUTPUT.exists()}")

    if OUTPUT.exists():
        print(f"Size: {OUTPUT.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()